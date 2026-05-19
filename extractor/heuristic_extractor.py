"""
Heuristic extractor — pulls structured chunks from raw documents
WITHOUT using LLMs. Output is a JSON-shaped dict that the LLM
validator will then verify.
"""
import re
import os
from typing import Dict, List, Optional, Any
from pathlib import Path

# ─── Document text loaders ─────────────────────────
def load_pdf_text(path: str) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ''
                text_parts.append(t)
                # Also pull tables (often where pricing lives)
                for tbl in page.extract_tables() or []:
                    for row in tbl:
                        text_parts.append(' | '.join(str(c or '') for c in row))
        return '\n'.join(text_parts)
    except Exception as e:
        print(f'  ⚠️ pdfplumber failed: {e}')
        try:
            from PyPDF2 import PdfReader
            r = PdfReader(path)
            return '\n'.join(p.extract_text() or '' for p in r.pages)
        except Exception as e2:
            print(f'  ⚠️ PyPDF2 also failed: {e2}')
            return ''

def load_xlsx_text(path: str) -> str:
    try:
        import pandas as pd
        xl = pd.ExcelFile(path)
        parts = []
        for sheet in xl.sheet_names:
            df = pd.read_excel(xl, sheet_name=sheet, header=None)
            parts.append(f'=== Sheet: {sheet} ===')
            parts.append(df.to_string(index=False, na_rep=''))
        return '\n'.join(parts)
    except Exception as e:
        print(f'  ⚠️ xlsx load failed: {e}')
        return ''

def load_docx_text(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                parts.append(' | '.join(c.text for c in row.cells))
        return '\n'.join(parts)
    except Exception as e:
        print(f'  ⚠️ docx load failed: {e}')
        return ''

def load_text(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == '.pdf': return load_pdf_text(path)
    if ext in ('.xlsx', '.xls'): return load_xlsx_text(path)
    if ext in ('.docx', '.doc'): return load_docx_text(path)
    if ext in ('.txt', '.csv'):
        return open(path, encoding='utf-8', errors='ignore').read()
    return ''

# ─── Heuristic field extractors ─────────────────────────
VENDOR_KEYWORDS = [
    'NTT Data', 'NTT DOCOMO', 'CDW', 'SHI', 'Equinix', 'Microsoft',
    'TrendMicro', 'Trend Micro', 'KnowBe4', 'ServiceNow', 'Quest',
    'Proquire', 'Honeywell', 'Copeland', 'Thrive', 'PC Connection',
    'Cisco', 'Palo Alto', 'Zscaler', 'CyberArk', 'Forescout',
    'NetApp', 'Oracle', 'VMware', 'Broadcom', 'Pure Storage', 'IBM'
]

CATEGORY_RULES = [
    ('Cybersecurity', ['cyber', 'security', 'firewall', 'endpoint', 'phishing',
                       'trend', 'cyberark', 'forescout', 'zscaler', 'knowbe4']),
    ('Network & Telecom', ['cisco', 'network', 'switch', 'router', 'wifi',
                           'wan', 'palo alto', 'meraki', 'wireless']),
    ('Hosting', ['vmware', 'oracle', 'netapp', 'colocation', 'hosting',
                 'server', 'storage', 'database', 'datacenter']),
    ('M365 & Power Platform', ['m365', 'office 365', 'microsoft 365', 'copilot',
                               'power bi', 'teams', 'sharepoint', 'visio']),
    ('Service Management (SNow)', ['servicenow', 'itsm', 'snow']),
    ('IdAM', ['active directory', 'identity', 'idam', 'quest', 'migration', 'odm']),
]

PROJECT_RULES = [
    ('Panasonic', ['panasonic', 'pas', 'pasap', 'pasj']),
    ('Idemia', ['idemia']),
    ('Tenneco', ['tenneco', 'lubrizol', 'copeland']),
]

def find_vendor(text: str, filename: str) -> str:
    blob = (text[:5000] + ' ' + filename).lower()
    for v in VENDOR_KEYWORDS:
        if v.lower() in blob:
            return v
    return 'Unknown'

def find_category(text: str, filename: str) -> str:
    blob = (text[:8000] + ' ' + filename).lower()
    scores = {}
    for cat, keywords in CATEGORY_RULES:
        score = sum(blob.count(kw.lower()) for kw in keywords)
        if score:
            scores[cat] = score
    return max(scores, key=scores.get) if scores else 'Cybersecurity'

def find_project(text: str, filename: str) -> str:
    blob = (text[:5000] + ' ' + filename).lower()
    for proj, keywords in PROJECT_RULES:
        if any(kw in blob for kw in keywords):
            return proj
    return 'Unknown'

def find_total_price(text: str) -> float:
    """Find the largest currency amount in the doc — usually the grand total."""
    patterns = [
        r'(?:grand\s*total|total\s*amount|total\s*price|sub\s*total|net\s*total)[:\s]*[\$€£]?\s*([\d,]+(?:\.\d{1,2})?)',
        r'[\$€£]\s*([\d,]+(?:\.\d{1,2})?)',
        r'USD\s*([\d,]+(?:\.\d{1,2})?)',
    ]
    candidates = []
    for pat in patterns:
        for m in re.finditer(pat, text, re.IGNORECASE):
            try:
                val = float(m.group(1).replace(',', ''))
                if 1000 < val < 100_000_000:
                    candidates.append(val)
            except (ValueError, IndexError):
                continue
    return max(candidates) if candidates else 0.0

def find_year(text: str, filename: str) -> int:
    blob = (filename + ' ' + text[:3000])
    matches = re.findall(r'\b(20\d{2})\b', blob)
    if matches:
        years = [int(m) for m in matches if 2020 <= int(m) <= 2030]
        if years:
            from collections import Counter
            return Counter(years).most_common(1)[0][0]
    return 2025

def find_quarter(text: str) -> Optional[str]:
    m = re.search(r'\bQ([1-4])\s*20\d{2}\b', text, re.IGNORECASE)
    if m: return f'Q{m.group(1)}'
    m = re.search(r'\b(jan|feb|mar)\w*\b', text[:2000], re.IGNORECASE)
    if m: return 'Q1'
    m = re.search(r'\b(apr|may|jun)\w*\b', text[:2000], re.IGNORECASE)
    if m: return 'Q2'
    m = re.search(r'\b(jul|aug|sep)\w*\b', text[:2000], re.IGNORECASE)
    if m: return 'Q3'
    m = re.search(r'\b(oct|nov|dec)\w*\b', text[:2000], re.IGNORECASE)
    if m: return 'Q4'
    return None

def find_quote_date(text: str) -> Optional[str]:
    """Returns ISO date string or None."""
    patterns = [
        (r'(?:date|dated|issued|valid\s+from)[:\s]+(\d{1,2}[\s/-]\w+[\s/-]\d{2,4})', '%d %B %Y'),
        (r'(?:date|dated)[:\s]+(\d{4}-\d{2}-\d{2})', '%Y-%m-%d'),
        (r'(?:date|dated)[:\s]+(\d{1,2}/\d{1,2}/\d{2,4})', '%m/%d/%Y'),
    ]
    from datetime import datetime
    for pat, fmt in patterns:
        m = re.search(pat, text[:3000], re.IGNORECASE)
        if m:
            try:
                d = datetime.strptime(m.group(1).strip(), fmt)
                return d.strftime('%Y-%m-%d')
            except ValueError:
                continue
    return None

def extract_line_items(text: str) -> List[Dict[str, Any]]:
    """
    Extract service line items as {name, sku, qty, unitPrice}.
    Looks for table-like rows separated by | or whitespace.
    """
    items = []
    lines = text.split('\n')
    
    # SKU pattern: alphanumeric with hyphens, often uppercase
    sku_re = re.compile(r'\b([A-Z][A-Z0-9]{2,}[-/][A-Z0-9-/]{2,})\b')
    # Currency pattern
    money_re = re.compile(r'[\$€£]?\s*([\d,]+\.\d{2})')
    # Quantity pattern (small integer)
    qty_re = re.compile(r'\b(\d{1,5})\b')
    
    for line in lines:
        if len(line.strip()) < 10:
            continue
        # Skip header lines
        if any(h in line.lower() for h in ['description', 'item no', 'part number',
                                             'sku', 'unit price', 'total']):
            continue
        
        skus = sku_re.findall(line)
        money = money_re.findall(line)
        
        if skus and money:
            sku = skus[0]
            unit_price = 0.0
            for m in money:
                try:
                    val = float(m.replace(',', ''))
                    if 0.01 < val < 100_000:
                        unit_price = val
                        break
                except ValueError:
                    pass
            
            qty = 1
            qty_matches = qty_re.findall(line)
            for q in qty_matches:
                try:
                    qv = int(q)
                    if 1 <= qv <= 10000 and str(qv) != sku.split('-')[-1]:
                        qty = qv
                        break
                except ValueError:
                    pass
            
            # Name = the remaining text
            name = re.sub(sku_re, '', line)
            name = re.sub(money_re, '', name)
            name = re.sub(r'[\d,]+\.\d{2}', '', name)
            name = re.sub(r'\s+', ' ', name).strip(' |')
            name = name[:80]
            
            if name and unit_price > 0:
                items.append({
                    'name': name,
                    'sku': sku,
                    'qty': qty,
                    'unitPrice': unit_price
                })
    
    return items[:30]  # cap

# ─── Main entry ─────────────────────────
def extract_chunks(filepath: str) -> Dict[str, Any]:
    """
    Returns a structured chunk dict that the LLM validator
    will then verify and enrich.
    """
    filename = os.path.basename(filepath)
    text = load_text(filepath)
    
    if not text or len(text) < 50:
        return {
            'file': filename,
            'extraction_status': 'failed',
            'reason': 'no_text_extracted'
        }
    
    chunk = {
        'file': filename,
        'extraction_status': 'heuristic',
        'raw_text_length': len(text),
        'raw_text_excerpt': text[:2000],  # for LLM context
        'vendor': find_vendor(text, filename),
        'project': find_project(text, filename),
        'category': find_category(text, filename),
        'price_total': find_total_price(text),
        'year': find_year(text, filename),
        'quarter': find_quarter(text),
        'quoteDate': find_quote_date(text),
        'services': extract_line_items(text),
    }
    return chunk
